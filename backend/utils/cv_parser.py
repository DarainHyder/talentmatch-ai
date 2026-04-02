"""
cv_parser.py
------------
Extracts raw text from uploaded CV files.
Supports: PDF (via PyPDF2) and DOCX (via python-docx).
Returns a cleaned plain-text string.
"""

import re
import io


def _clean_text(text: str) -> str:
    """
    Remove excessive whitespace, non-printable characters,
    and normalise newlines while keeping sentence structure.
    """
    # Replace tabs and non-breaking spaces with regular space
    text = text.replace("\t", " ").replace("\xa0", " ")
    # Collapse multiple blank lines into a single newline
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces into one
    text = re.sub(r" {2,}", " ", text)
    # Remove non-ASCII characters that are just noise
    text = re.sub(r"[^\x00-\x7F]+", " ", text)
    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    # Drop completely empty lines at start/end; keep internal structure
    text = "\n".join(lines).strip()
    return text


def parse_cv(file) -> str:
    """
    Parse a CV file object and return extracted plain text.

    Parameters
    ----------
    file : file-like object (from Flask request.files) OR a file path string.
           The filename attribute (or the path extension) determines the parser.

    Returns
    -------
    str : cleaned raw text from the CV
    
    Raises
    ------
    ValueError : if the file type is not PDF or DOCX
    RuntimeError : if extraction fails
    """
    # Determine file type
    filename = ""
    if hasattr(file, "filename"):
        filename = file.filename.lower()
    elif isinstance(file, str):
        filename = file.lower()

    if filename.endswith(".pdf"):
        return _parse_pdf(file)
    elif filename.endswith(".docx"):
        return _parse_docx(file)
    else:
        raise ValueError(
            f"Unsupported file type: '{filename}'. "
            "Only PDF and DOCX files are supported."
        )


def _parse_pdf(file) -> str:
    """Extract text from a PDF using PyPDF2."""
    try:
        import PyPDF2

        # Accept both file-like objects and file paths
        if isinstance(file, str):
            f = open(file, "rb")
            close_after = True
        else:
            # Flask file objects: read into BytesIO so PyPDF2 can seek
            content = file.read()
            f = io.BytesIO(content)
            close_after = False

        reader = PyPDF2.PdfReader(f)
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text)

        if close_after:
            f.close()

        raw = "\n".join(pages)
        if not raw.strip():
            raise RuntimeError("PDF appears to be empty or image-based (no extractable text).")

        return _clean_text(raw)

    except ImportError:
        raise RuntimeError("PyPDF2 is not installed. Run: pip install PyPDF2")
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF: {e}")


def _parse_docx(file) -> str:
    """Extract text from a DOCX using python-docx."""
    try:
        from docx import Document

        # Accept both file-like objects and file paths
        if isinstance(file, str):
            doc = Document(file)
        else:
            content = file.read()
            doc = Document(io.BytesIO(content))

        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        raw = "\n".join(paragraphs)

        if not raw.strip():
            raise RuntimeError("DOCX appears to be empty (no extractable paragraphs).")

        return _clean_text(raw)

    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX: {e}")
